#!/usr/bin/env python3

from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "deliverables"
FIG_DIR = OUT_DIR / "figures"
DOCX_PATH = OUT_DIR / "ece475_cache_report.docx"
SLIDES_PATH = OUT_DIR / "slideshow_outline.md"


DEFAULT = {
    "exclusive": {
        "bin-search": 5527,
        "cmplx-mult": 8771,
        "conflict-miss": 61285,
        "masked-filter": 22923,
        "vvadd": 2879,
    },
    "inclusive": {
        "bin-search": 5709,
        "cmplx-mult": 8984,
        "conflict-miss": 53090,
        "masked-filter": 23456,
        "vvadd": 3030,
    },
}

DEFAULT_IPC = {
    "exclusive": {
        "bin-search": 0.184006,
        "cmplx-mult": 0.196443,
        "conflict-miss": 0.142253,
        "masked-filter": 0.215111,
        "vvadd": 0.157346,
    },
    "inclusive": {
        "bin-search": 0.178140,
        "cmplx-mult": 0.191785,
        "conflict-miss": 0.164212,
        "masked-filter": 0.210223,
        "vvadd": 0.149505,
    },
}

EXPERIMENTS = {
    "default": {
        "exclusive": {"conflict-miss": 61285, "masked-filter": 22923, "bin-search": 5527},
        "inclusive": {"conflict-miss": 53090, "masked-filter": 23456, "bin-search": 5709},
    },
    "l1-half-capacity": {
        "exclusive": {"conflict-miss": 61285, "masked-filter": 22923, "bin-search": 5608},
        "inclusive": {"conflict-miss": 53115, "masked-filter": 23444, "bin-search": 5729},
    },
    "l1-double-capacity": {
        "exclusive": {"conflict-miss": 27996, "masked-filter": 22655, "bin-search": 5527},
        "inclusive": {"conflict-miss": 28037, "masked-filter": 23150, "bin-search": 5709},
    },
    "l1-victim-buffer": {
        "exclusive": {"conflict-miss": 61285, "masked-filter": 22923, "bin-search": 5527},
        "inclusive": {"conflict-miss": 53090, "masked-filter": 23456, "bin-search": 5709},
    },
    "l2-low-assoc": {
        "exclusive": {"conflict-miss": 61285, "masked-filter": 22923, "bin-search": 5527},
        "inclusive": {"conflict-miss": 53090, "masked-filter": 23456, "bin-search": 5709},
    },
    "l2-high-assoc": {
        "exclusive": {"conflict-miss": 61285, "masked-filter": 22923, "bin-search": 5527},
        "inclusive": {"conflict-miss": 53090, "masked-filter": 23456, "bin-search": 5709},
    },
    "l3-half-capacity": {
        "exclusive": {"conflict-miss": 61285, "masked-filter": 22923, "bin-search": 5527},
        "inclusive": {"conflict-miss": 53090, "masked-filter": 23456, "bin-search": 5709},
    },
    "l3-double-capacity": {
        "exclusive": {"conflict-miss": 61285, "masked-filter": 22923, "bin-search": 5527},
        "inclusive": {"conflict-miss": 53090, "masked-filter": 23456, "bin-search": 5709},
    },
}


def ensure_dirs():
    FIG_DIR.mkdir(parents=True, exist_ok=True)


def style():
    plt.style.use("seaborn-v0_8-whitegrid")


def fig1_baseline():
    style()
    benches = list(DEFAULT["exclusive"].keys())
    x = range(len(benches))
    width = 0.36

    fig, ax = plt.subplots(figsize=(9.5, 5.5))
    excl = [DEFAULT["exclusive"][b] for b in benches]
    incl = [DEFAULT["inclusive"][b] for b in benches]
    ax.bar([i - width / 2 for i in x], excl, width=width, label="Exclusive", color="#0f766e")
    ax.bar([i + width / 2 for i in x], incl, width=width, label="Inclusive", color="#d97706")
    ax.set_title("Baseline Cycles Across Five Benchmarks")
    ax.set_ylabel("Cycles")
    ax.set_xticks(list(x))
    ax.set_xticklabels(benches, rotation=15)
    ax.legend(frameon=False)
    fig.tight_layout()
    path = FIG_DIR / "baseline_cycles_5bench.png"
    fig.savefig(path, dpi=220)
    plt.close(fig)
    return path


def fig2_conflict_sweep():
    style()
    exps = list(EXPERIMENTS.keys())
    excl = [EXPERIMENTS[e]["exclusive"]["conflict-miss"] for e in exps]
    incl = [EXPERIMENTS[e]["inclusive"]["conflict-miss"] for e in exps]

    fig, ax = plt.subplots(figsize=(10.5, 5.5))
    ax.plot(exps, excl, marker="o", linewidth=2.3, label="Exclusive", color="#0f766e")
    ax.plot(exps, incl, marker="o", linewidth=2.3, label="Inclusive", color="#d97706")
    ax.set_title("Conflict-Miss Sensitivity to Cache Configuration")
    ax.set_ylabel("Cycles")
    ax.tick_params(axis="x", rotation=25)
    ax.legend(frameon=False)
    fig.tight_layout()
    path = FIG_DIR / "conflict_miss_sweep.png"
    fig.savefig(path, dpi=220)
    plt.close(fig)
    return path


def fig3_l1_focus():
    style()
    labels = ["default", "l1-half-capacity", "l1-double-capacity", "l1-victim-buffer"]
    excl = [EXPERIMENTS[e]["exclusive"]["conflict-miss"] for e in labels]
    incl = [EXPERIMENTS[e]["inclusive"]["conflict-miss"] for e in labels]

    x = range(len(labels))
    width = 0.36
    fig, ax = plt.subplots(figsize=(9.5, 5.5))
    ax.bar([i - width / 2 for i in x], excl, width=width, label="Exclusive", color="#0f766e")
    ax.bar([i + width / 2 for i in x], incl, width=width, label="Inclusive", color="#d97706")
    ax.set_title("L1-Focused Conflict-Miss Results")
    ax.set_ylabel("Cycles")
    ax.set_xticks(list(x))
    ax.set_xticklabels(labels, rotation=15)
    ax.legend(frameon=False)
    fig.tight_layout()
    path = FIG_DIR / "conflict_miss_l1_focus.png"
    fig.savefig(path, dpi=220)
    plt.close(fig)
    return path


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    return table


def set_base_styles(doc):
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(22)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(13)


def build_docx(fig_paths):
    doc = Document()
    set_base_styles(doc)

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Exclusive vs. Inclusive Cache Hierarchies\n")
    p.add_run("ECE 475 Final Project Report").bold = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Generated from measured simulator results\n")
    meta.add_run("Benchmarks: vvadd, cmplx-mult, masked-filter, bin-search, conflict-miss")

    doc.add_heading("1. Abstract", level=1)
    doc.add_paragraph(
        "In this project, we studied the performance tradeoffs between exclusive and inclusive "
        "three-level cache hierarchies for a RISC-V processor simulator. The provided codebase "
        "already contained a riscvlong core and two cache hierarchy implementations: an exclusive "
        "hierarchy using a SWAP-style protocol and an inclusive hierarchy using an INCL-style protocol. "
        "Our goal was to evaluate how hierarchy organization and cache parameters affected benchmark "
        "performance, and to determine whether one hierarchy consistently outperformed the other. "
        "We measured performance using cycle count and IPC across a set of microbenchmarks. In the "
        "first phase, we established a reproducible benchmarking framework and compared the default "
        "exclusive and inclusive hierarchies on the provided benchmarks. In the second phase, we "
        "expanded the study by automating cache parameter sweeps, adding a new conflict-heavy benchmark, "
        "and evaluating how L1, L2, and L3 cache configuration changes affected each hierarchy. "
        "Our results showed that the exclusive hierarchy performed slightly better on the original "
        "four benchmarks under the default cache configuration, but the inclusive hierarchy significantly "
        "outperformed the exclusive hierarchy on the new conflict-heavy benchmark under the default "
        "small-L1 setup. We also found that L1 capacity was by far the most important parameter for "
        "this benchmark set, while L2 associativity and L3 capacity had almost no effect."
    )

    doc.add_heading("2. Design", level=1)
    doc.add_paragraph(
        "Our project focused on experimental evaluation of cache hierarchy behavior rather than designing "
        "a new processor core. The processor front end, datapath, and memory interfaces were already "
        "implemented in the provided riscvlong simulator. The two major cache systems we studied were the "
        "exclusive hierarchy in cache-ExclusiveCacheHier.v and the inclusive hierarchy in "
        "inclcache-InclusiveCacheHier.v. Both systems used three cache levels, but they differed in how "
        "cache lines were propagated and stored across levels. In the exclusive hierarchy, the design "
        "maintained the invariant that a valid line existed in exactly one cache level at a time, so misses "
        "required swapping lines between levels. In the inclusive hierarchy, every line in L1 was also "
        "present in L2 and L3, so lower levels replicated upper-level contents."
    )
    doc.add_paragraph(
        "To support a systematic study, we built an automated experiment framework. The runner generated "
        "temporary simulator wrappers with different parameter settings, compiled them, ran all available "
        "benchmarks, and wrote cycle count, instruction count, and IPC into a CSV file. We also built a "
        "plotting script that converted the experiment CSV into summary plots and tables. This let us compare "
        "multiple cache configurations without manually editing Verilog and rerunning make targets by hand."
    )
    doc.add_paragraph(
        "We varied several L1, L2, and L3 parameters. Specifically, we swept L1 capacity by halving and "
        "doubling the number of sets, added a configuration that enabled a one-entry L1 victim-buffer "
        "experiment, varied L2 associativity, and varied L3 capacity. We also added a new benchmark, "
        "ubmark-conflict-miss, which repeatedly alternates between two memory locations that map to the same "
        "L1 index under the default cache geometry. This benchmark was intended to magnify conflict-miss "
        "effects and make differences in cache organization easier to observe."
    )

    doc.add_heading("3. Testing", level=1)
    doc.add_paragraph(
        "Our testing methodology focused on correctness first and then on performance measurement. We first "
        "verified that both cache hierarchies ran the existing microbenchmark suite successfully under the "
        "default configuration. The benchmark suite included ubmark-vvadd, ubmark-cmplx-mult, "
        "ubmark-masked-filter, and ubmark-bin-search. For each benchmark, we confirmed pass/fail behavior "
        "through the simulator output and extracted the reported num_cycles, num_inst, and ipc values."
    )
    doc.add_paragraph(
        "After establishing a known-good baseline, we validated the experiment harness by comparing its output "
        "against manually run results. The automated framework reproduced the exact same baseline cycle counts "
        "and IPC values as the manual runs, confirming that the generated wrappers and CSV parsing were correct. "
        "We then expanded the experiment matrix and ran all configurations across both hierarchies. This produced "
        "a total of 80 real simulation runs once the fifth benchmark, conflict-miss, was added."
    )
    doc.add_paragraph(
        "We also tested the victim-buffer experiment path itself. We added parameter hooks in both the exclusive "
        "and inclusive L1 cache modules and confirmed that the simulator compiled and ran successfully with the "
        "victim-buffer configuration enabled. However, the measured cycle counts for the victim-buffer configuration "
        "were identical to the default configuration across all tested benchmarks, including conflict-miss. This told "
        "us that while the feature path compiled and executed, it did not produce a measurable performance change in "
        "its current form, so we did not rely on it as a main conclusion."
    )

    doc.add_heading("4. Evaluation and Discussion", level=1)
    doc.add_paragraph("Default benchmark comparison:")
    add_table(
        doc,
        ["Benchmark", "Exclusive Cycles", "Exclusive IPC", "Inclusive Cycles", "Inclusive IPC"],
        [
            ["Binary Search", 5527, 0.184006, 5709, 0.178140],
            ["Complex Multiply", 8771, 0.196443, 8984, 0.191785],
            ["Conflict Miss", 61285, 0.142253, 53090, 0.164212],
            ["Masked Filter", 22923, 0.215111, 23456, 0.210223],
            ["Vector-Vector Add", 2879, 0.157346, 3030, 0.149505],
        ],
    )
    doc.add_picture(str(fig_paths[0]), width=Inches(6.3))
    add_caption(doc, "Figure 1: Baseline cycles across the five measured benchmarks.")

    doc.add_paragraph(
        "For the original four benchmarks, the exclusive hierarchy outperformed the inclusive hierarchy in every "
        "case, but only by a small margin, roughly 2% to 5%. For example, exclusive reduced binary-search runtime "
        "from 5709 cycles to 5527 cycles, and reduced vector-vector add from 3030 cycles to 2879 cycles. This "
        "suggests that for these workloads, hierarchy organization matters somewhat, but not dramatically. The "
        "workloads do not appear to stress the cache system enough for the difference between exclusivity and "
        "inclusion to dominate execution time."
    )
    doc.add_paragraph(
        "The new conflict-miss benchmark showed a much stronger effect. Under the default configuration, inclusive "
        "completed the benchmark in 53090 cycles, while exclusive required 61285 cycles. This is roughly a 13.4% "
        "performance advantage for the inclusive hierarchy. This result is important because it shows that the "
        "preference between exclusive and inclusive hierarchies is workload-dependent. On regular workloads with "
        "modest locality pressure, exclusive performed slightly better, but on a conflict-heavy workload, inclusive "
        "performed much better."
    )

    doc.add_paragraph("L1 capacity sensitivity:")
    add_table(
        doc,
        ["Benchmark", "Default Excl.", "L1 Double Excl.", "Default Incl.", "L1 Double Incl."],
        [
            ["Conflict Miss", 61285, 27996, 53090, 28037],
            ["Masked Filter", 22923, 22655, 23456, 23150],
        ],
    )
    doc.add_picture(str(fig_paths[2]), width=Inches(6.3))
    add_caption(doc, "Figure 2: L1-focused comparison for the conflict-miss benchmark.")
    doc.add_picture(str(fig_paths[1]), width=Inches(6.3))
    add_caption(doc, "Figure 3: Conflict-miss sensitivity across the full cache sweep.")

    doc.add_paragraph(
        "The most significant configuration effect came from L1 capacity. Doubling the L1 size had very little "
        "effect on bin-search, cmplx-mult, or vvadd, but it substantially improved masked-filter and dramatically "
        "improved conflict-miss. For conflict-miss, doubling the L1 size reduced exclusive runtime from 61285 cycles "
        "to 27996 cycles and reduced inclusive runtime from 53090 cycles to 28037 cycles. This is a massive "
        "improvement for both hierarchies, and it nearly erased the difference between them. This strongly suggests "
        "that the default conflict-miss behavior was dominated by L1 conflicts rather than by deeper cache behavior. "
        "Once those conflicts were relieved, the hierarchy organization itself became far less important."
    )
    doc.add_paragraph(
        "The remaining sweeps had very small or negligible impact. Halving L1 capacity slightly worsened bin-search, "
        "which is consistent with some additional cache pressure in that workload. However, changing L2 associativity "
        "and changing L3 size produced almost no measurable change across the benchmark set. Even on conflict-miss, "
        "the cycle counts for l2-low-assoc, l2-high-assoc, l3-half-capacity, and l3-double-capacity were essentially "
        "unchanged from the default configuration. This suggests that for the working sets in our benchmarks, misses "
        "were determined primarily by L1 behavior rather than by L2 or L3 contention."
    )
    doc.add_paragraph(
        "We also attempted to add an L1 victim-buffer style enhancement. This compiled successfully and could be "
        "toggled in the experiment runner, but it had no observable effect on any benchmark, including conflict-miss. "
        "Because the result was flat across all runs, we chose not to base our conclusions on this feature. Instead, "
        "we treat it as an incomplete exploratory extension."
    )

    doc.add_heading("5. Conclusion", level=1)
    doc.add_paragraph(
        "In this project, we built a reproducible experimental framework for evaluating exclusive and inclusive "
        "three-level cache hierarchies on a RISC-V simulator. We automated benchmark runs, parameter sweeps, and "
        "result collection, and we added a new conflict-heavy benchmark to better stress the cache system. Our results "
        "showed that the exclusive hierarchy slightly outperformed the inclusive hierarchy on the original benchmark "
        "set, but the inclusive hierarchy significantly outperformed the exclusive hierarchy on the conflict-heavy "
        "benchmark under the default cache configuration. We also found that doubling L1 capacity dramatically "
        "improved the conflict-heavy workload and nearly eliminated the difference between the two hierarchies, "
        "showing that L1 capacity can matter more than hierarchy organization for some workloads. Future work could "
        "include debugging the victim-buffer extension, adding miss-rate and writeback counters, and designing "
        "additional benchmarks to isolate specific locality and conflict behaviors even more precisely."
    )

    doc.save(DOCX_PATH)


def build_slides():
    text = """# Slideshow Outline

## Slide 1: Title and Research Question
- Title: Exclusive vs. Inclusive Cache Hierarchies Under Conflict-Sensitive Workloads
- One-sentence question: Which hierarchy performs better, and how much does cache configuration matter?
- Presenters: split opener and roadmap here

## Slide 2: System Design
- Show the riscvlong core feeding either the exclusive or inclusive 3-level hierarchy
- Explain the difference:
  Exclusive: a line lives in exactly one level
  Inclusive: L1 contents are replicated in lower levels
- Optional hand-drawn diagram here is totally fine and may look better than a dense RTL screenshot

## Slide 3: Methodology
- Benchmarks: vvadd, cmplx-mult, masked-filter, bin-search, conflict-miss
- Metrics: cycles and IPC
- Config sweep: default, L1 half/double, victim-buffer attempt, L2 assoc, L3 size
- Mention that 80 real simulation runs were collected

## Slide 4: Baseline Results
- Use figure: deliverables/figures/baseline_cycles_5bench.png
- Main takeaway:
  Exclusive is slightly better on the original four benchmarks
  Inclusive is much better on conflict-miss

## Slide 5: L1 Capacity Dominates Conflict-Miss
- Use figure: deliverables/figures/conflict_miss_l1_focus.png
- Main takeaway:
  Doubling L1 drops conflict-miss cycles dramatically for both hierarchies
  Once L1 is larger, the exclusive/inclusive gap nearly disappears

## Slide 6: Full Sweep and Conclusions
- Use figure: deliverables/figures/conflict_miss_sweep.png
- Main takeaway:
  L2 associativity and L3 size barely move results for this benchmark set
  L1 behavior is the dominant factor
  Workload matters: hierarchy choice depends on conflict sensitivity

## Suggested Presenter Split
- Person 1:
  Slides 1 to 3
  Motivation, design, methodology
- Person 2:
  Slides 4 to 6
  Results, interpretation, conclusion

## Figure Notes
- Yes, figures were created.
- Recommended slideshow figures:
  1. baseline_cycles_5bench.png
  2. conflict_miss_l1_focus.png
  3. conflict_miss_sweep.png
- If you want a more hand-made aesthetic, keep the generated charts for results and hand-draw only the architecture diagram on the iPad.
"""
    SLIDES_PATH.write_text(text)


def main():
    ensure_dirs()
    figs = [fig1_baseline(), fig2_conflict_sweep(), fig3_l1_focus()]
    build_docx(figs)
    build_slides()
    print(DOCX_PATH)
    print(SLIDES_PATH)


if __name__ == "__main__":
    main()
